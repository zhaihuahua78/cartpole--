"""生成 Word 版实验报告（从空白文档创建）"""
import os, sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC = r"D:\大二下\AI赋能自动控制原理\AI赋能自动控制原理\AI赋能自动控制原理\task1_cartpole"
os.chdir(SRC)

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_title(text, level=1):
    """Add heading with proper Chinese formatting"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, size=12, align=None, first_indent=True):
    p = doc.add_paragraph()
    if first_indent and len(text) > 5:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if align:
        p.alignment = align
    return p

def add_image(path, width_cm=12):
    full = os.path.join(SRC, path)
    if os.path.exists(full):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(full, width=Cm(width_cm))
    return p

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.italic = True
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════
# Title
# ═══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("任务1实验报告")
run.font.size = Pt(22)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("基于强化学习的倒立摆控制")
run.font.size = Pt(16)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("课程：AI赋能自动控制原理    日期：2026年7月")
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ════════════════════ 1 ════════════════════
add_title("1 问题分析", level=1)
add_title("1.1 Cart Pole系统与强化学习环境", level=2)
add_title("1.1.1 状态空间与动作空间", level=3)

add_para("Cart Pole（小车-倒立摆）是自动控制原理中经典的非线性系统控制案例。系统由一辆可在轨道上水平移动的小车和一根铰接于小车上的摆杆组成，控制目标是通过对小车施加水平推力（左移或右移），使摆杆始终保持直立状态。")
add_para("本实验使用OpenAI Gymnasium的CartPole-v1物理引擎，并按照任务描述的要求进行了定制：终止角度设为15°（原默认12°），失败时施加-10惩罚。状态空间为4维连续向量，动作空间为离散二值：动作0表示向左推，动作1表示向右推。环境的终止条件有三个：（1）摆杆倾斜超过15°，视为失去平衡；（2）小车偏离中心超过2.4m，视为滑出轨道；（3）达到500步且摆杆未倒下，视为成功。")

add_table(
    ["序号", "符号", "含义", "单位"],
    [["0", "x", "小车位置", "m"],
     ["1", "ẋ", "小车速度", "m/s"],
     ["2", "θ", "摆杆角度", "rad"],
     ["3", "θ̇", "摆杆角速度", "rad/s"]]
)

add_title("1.2 算法与奖励函数设计", level=2)
add_title("1.2.1 算法选择及基本原理", level=3)

add_para("本实验选用两种主流的深度强化学习算法进行对比研究：")
add_para("PPO（Proximal Policy Optimization）：基于策略梯度的算法，通过裁剪机制（clip_range=0.2）限制策略更新幅度，防止单次更新过大导致性能崩溃。PPO的核心优势在于训练稳定性高、超参数鲁棒性好。关键超参数：学习率3×10⁻⁴、n_steps=2,048、GAE λ=0.95、折扣因子γ=0.99。")
add_para("A2C（Advantage Actor-Critic）：同步优势演员-评论家算法，同时学习策略网络和价值网络，利用优势函数降低策略梯度方差。经7种配置的系统扫描，最优超参数为：n_steps=64、学习率5×10⁻⁴、熵系数0.02。")
add_para("传统控制方法：随机策略、规则控制器、LQR控制器作为性能基线。")
add_para("奖励函数设计。严格遵循任务描述：每步成功保持平衡获得+1正向奖励，失败（角度或位置超限）获得-10负向惩罚。")

# ════════════════════ 2 ════════════════════
add_title("2 实验过程", level=1)
add_title("2.1 环境构建与参数设置", level=2)
add_para("本实验使用自定义的CartPole环境（src/task_env.py），基于OpenAI Gymnasium的物理引擎实现，终止角度设为15°、失败惩罚-10。环境通过DummyVecEnv进行向量化，配合Monitor包装器记录每回合的奖励与步数。训练统一配置：训练步数200,000步、随机种子42、CPU设备。")

add_table(
    ["参数", "PPO", "A2C（优化后）"],
    [["学习率", "3×10⁻⁴", "5×10⁻⁴"],
     ["n_steps", "2,048", "64"],
     ["batch_size", "64", "--"],
     ["n_epochs", "10", "--"],
     ["gamma", "0.99", "0.99"],
     ["GAE λ", "0.95", "0.95"],
     ["clip_range", "0.2", "--"],
     ["ent_coef", "0.0", "0.02"]]
)

add_title("2.2 智能体训练过程", level=2)
add_title("2.2.1 PPO训练过程", level=3)
add_para("PPO在200,000步内完美求解自定义环境。训练过程分为三个阶段：探索期（0-200回合）奖励10-40；收敛期（200-280回合）奖励从40快速攀升至500；稳定求解期（280回合后）持续满分，后100回合均值500.0。20回合评估全部满分（500.0±0.0），达标率100%。")
add_image("results_task/fig1_ppo_training.png", width_cm=10)
add_caption("图1 PPO训练奖励曲线")

add_title("2.2.2 A2C训练过程", level=3)
add_para("A2C经超参数优化后（n_steps=64, lr=5e-4, ent_coef=0.02）在200,000步内接近收敛。训练过程：缓慢上升期（0-200回合）均值29→227；波动求解期（200-500回合）多次满分但波动；稳定提升期（500回合后）后100回合均值464.5。20回合评估均值493.4±20.8，达标率90%（18/20）。")
add_image("results_task/fig2_a2c_training.png", width_cm=12)
add_caption("图2 A2C训练奖励曲线")

add_title("2.3 关键参数验证", level=2)
add_title("2.3.1 PPO关键参数验证", level=3)
add_para("PPO的核心超参数为学习率。实验对比了三种学习率：1e-4（较慢，后50均值约180）、3e-4（适中，后50均值约350）、1e-3（最快但可能发散，后50均值约280）。3×10⁻⁴是平衡收敛速度与稳定性的最优选择。")

add_title("2.3.2 A2C关键参数验证", level=3)
add_para("A2C对超参数更敏感。对7种配置进行了系统扫描，关键结论：n_steps=64最优（后100均值433.5，求解67次），n_steps过小（5）样本利用率低（后100均值41.2），过大（256/512）更新频率不足。综合后100均值和求解次数，n_steps=64, lr=5e-4为最优配置。")

add_image("results_task/fig_s1_ppo_lr.png", width_cm=8)
add_caption("图3 PPO学习率对比")
add_image("results_task/fig_s2_a2c_nsteps.png", width_cm=8)
add_caption("图4 A2C n_steps对比")
add_image("results_task/fig_s3_a2c_default_vs_optimal.png", width_cm=8)
add_caption("图5 默认vs最优A2C")

add_title("2.4 扩展对比（传统控制方法）", level=2)
add_para("从设计思路、是否需要训练、控制流程三个维度对比传统控制方法：")
add_para("随机策略：每步50%随机动作，无需训练，作为性能下界。LQR控制器：在平衡点线性化后求解Riccati方程获得最优反馈增益K。规则控制器：基于启发式规则——摆杆向右倾斜则向右推、向左倾斜则向左推，无需建模和训练。")

add_table(
    ["方法", "是否需要训练", "训练时间", "设计复杂度"],
    [["随机策略", "否", "0", "无"],
     ["LQR控制器", "否", "0", "高"],
     ["规则控制器", "否", "0", "低"],
     ["A2C", "是", "约3分钟", "中"],
     ["PPO", "是", "约2分钟", "低"]]
)

# ════════════════════ 3 ════════════════════
add_title("3 实验结果及分析", level=1)
add_title("3.1 训练过程分析", level=2)
add_title("3.1.1 PPO训练过程分析", level=3)
add_para("PPO的奖励曲线呈典型S形增长。探索期（0-200回合）奖励10-40；收敛期（200-280回合）奖励从40陡峭攀升至500，每10回合平均提升约80分；求解期（280回合后）奖励密集分布在满分附近。从多指标面板分析：滚动均值从约50分匀速增长至500分后严格稳定；成功率在200-300回合从0%跃升至60%以上；回合步数方差从±200步缩小至±50步以内。")

add_title("3.1.2 A2C训练过程分析", level=3)
add_para("A2C的奖励曲线呈多尖峰模式，非平滑S形。后100回合均值达464.5，满分尖峰频率在后期增加。滚动均值呈阶梯状增长（200-400回合约150分、400-550回合约250分、550-700回合约380分）。成功率在大部分训练时间内接近0%，仅后200回合出现正值且不超过30%。")

add_image("results_task/fig6_metrics_panel.png", width_cm=14)
add_caption("图6 训练过程多指标对比面板")

add_title("3.2 训练前后控制效果", level=2)
add_title("3.2.1 PPO训练前后对比", level=3)
add_para("训练前（随机策略）中位数约15分。训练后（PPO）箱体压缩为单值500分（零方差），提升超过31倍。标准差从6.8降至0，策略对不同的初始条件具有完全的鲁棒性。")

add_title("3.2.2 A2C训练前后对比", level=3)
add_para("训练后A2C中位数达500分，均值493.4分，提升约31倍。与PPO的关键区别在于存在3个异常值（60-130分），标准差20.8远高于PPO的0。策略能力与PPO接近但稳定性差距明显。")

add_image("results_task/fig5_before_after.png", width_cm=9)
add_caption("图7 训练前后控制效果对比")

add_title("3.3 对比结果", level=2)
add_para("五种控制方法平均奖励排序为：PPO（500.0）> A2C（493.4）≫ 规则控制器（215.8）> LQR（44.8）> 随机策略（16.1）。PPO与A2C的均值差距仅1.4%，但PPO标准差为0而A2C为20.8——稳定性差距显著。规则控制器是传统方法天花板，但仅达到PPO性能的43.2%。LQR仅比随机策略高约2.8倍，线性化模型在±15°的非线性范围内基本失效。")

add_image("results_task/fig3_method_comparison.png", width_cm=10)
add_caption("图8 五种控制方法平均奖励对比")
add_image("results_task/fig4_training_comparison.png", width_cm=12)
add_caption("图9 PPO与A2C训练过程对比")

add_table(
    ["方法", "平均奖励", "标准差", "达标率", "训练步数", "是否收敛"],
    [["随机策略", "16.1", "6.8", "0%", "0", "否"],
     ["LQR", "44.8", "7.4", "0%", "0", "否（模型失配）"],
     ["规则控制", "215.8", "29.1", "0%", "0", "否"],
     ["A2C（优化）", "493.4", "20.8", "90%", "200,000", "接近收敛"],
     ["PPO", "500.0", "0.0", "100%", "200,000", "完全收敛"]]
)

# ════════════════════ 4 ════════════════════
add_title("4 总结体会", level=1)
add_title("4.1 算法效果总结", level=2)
add_para("本实验严格按任务描述实现了自定义CartPole环境（15°终止角、-10失败惩罚），系统对比了五种控制方法。PPO以满分500、零标准差完美解决问题，是唯一兼具高能力与高稳定性的算法。优化后的A2C（n_steps=64, lr=5e-4）评估均值493.4、达标率90%，性能接近PPO。传统控制方法中，规则控制器（215.8分）大幅超越LQR（44.8分）和随机策略（16.1分），证明了在强非线性系统中启发式规则的有效性。")

add_title("4.2 存在的问题与可改进方向", level=2)
add_para("算法覆盖度不足：仅对比了PPO和A2C，可引入SAC、TD3等更先进的算法。A2C稳定性不足：仍有10%评估回合失败。超参数搜索不充分：仅覆盖了n_steps和学习率。泛化性验证缺失：未测试物理参数变化时的鲁棒性。可引入课程学习或模仿学习加速训练。")

add_title("4.3 个人收获", level=2)
add_para("掌握了PPO和A2C的核心原理和实现差异——PPO的裁剪机制如何保证稳定性，A2C的演员-评论家架构如何降低方差。理解了超参数调优的决定性影响——A2C从默认参数（后100均值41.2）到优化参数（后100均值464.5）的性能飞跃是最佳证明。认识到传统控制方法与RL方法的本质差异：前者依赖精确建模或人工规则，后者通过与环境试错交互自主学习策略。")

# ════════════════════ 5 ════════════════════
add_title("5 附录", level=1)
add_para("本项目代码已开源至GitHub：")
add_para("https://github.com/zhaihuahua78/cartpole--")
doc.add_paragraph()
add_para("报告完成日期：2026年7月", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
output = os.path.join(SRC, "task1_report.docx")
doc.save(output)
print(f"Word report saved to: {output}")
print(f"File size: {os.path.getsize(output)} bytes")
